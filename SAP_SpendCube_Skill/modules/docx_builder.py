"""
SAP Procurement Spend Cube - Word Document Builder
=====================================================
Creates comprehensive DOCX reports with all rules, mappings, and metrics.
Supports full-dataset and single-company modes.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import (DOCX_STYLES, SAP_TABLES, RSEG_JOIN_SEQUENCE, SPEND_CLASSIFICATION,
                    SHKZG_RULES, COMPANY_CODES, MAVERICK_TIERS, THREE_WAY_MATCH,
                    GL_ACCOUNT_DESCRIPTIONS, SAVINGS_ESTIMATES,
                    safe_vendor_name, get_analysis_date)

DATE = get_analysis_date()


# ============================================================================
# PRIMITIVES
# ============================================================================

def create_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = DOCX_STYLES['normal_font']
    style.font.size = Pt(DOCX_STYLES['normal_size_pt'])
    for i in range(1, 4):
        doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(*DOCX_STYLES['heading_color_rgb'])
    return doc

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = DOCX_STYLES['table_style']
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: run.bold = True; run.font.size = Pt(DOCX_STYLES['table_header_font_size_pt'])
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            cell = table.rows[ri + 1].cells[ci]; cell.text = str(ct)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(DOCX_STYLES['table_body_font_size_pt'])
    doc.add_paragraph()
    return table

def add_code_block(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = DOCX_STYLES['code_font']
    run.font.size = Pt(DOCX_STYLES['code_size_pt'])
    run.font.color.rgb = RGBColor(*DOCX_STYLES['code_color_rgb'])

def add_cover_page(doc, title, subtitle, sub2, info_text):
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title); run.font.size = Pt(DOCX_STYLES['cover_title_size_pt'])
    run.bold = True; run.font.color.rgb = RGBColor(*DOCX_STYLES['heading_color_rgb'])

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle); run.font.size = Pt(DOCX_STYLES['cover_subtitle_size_pt'])
    run.font.color.rgb = RGBColor(255, 140, 0)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sub2); run.font.size = Pt(DOCX_STYLES['cover_sub2_size_pt'])
    run.font.color.rgb = RGBColor(*DOCX_STYLES['heading_color_rgb'])

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(info_text); run.font.size = Pt(DOCX_STYLES['cover_info_size_pt'])
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

def add_toc(doc, items):
    doc.add_heading('Table of Contents', level=1)
    for item in items:
        p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(2)
        p.runs[0].font.size = Pt(DOCX_STYLES['toc_item_size_pt'])
    doc.add_page_break()


# ============================================================================
# STANDARD SECTIONS
# ============================================================================

def section_executive_summary(doc, m, title="Full Dataset", code="All"):
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(f'Comprehensive procurement spend analysis for {title} (Company Code: {code}).')
    rows = [
        ['Total Invoice Line Items', f'{m["total_records"]:,}'],
        ['Total Gross Spend', f'EUR {m["total_gross"]/1e6:.1f}M'],
        ['Total Net Spend', f'EUR {m["total_net"]/1e6:.1f}M'],
        ['Total Credits', f'EUR {m["total_credit"]/1e6:.1f}M'],
        ['Direct Spend', f'EUR {m["direct_spend"]/1e6:.1f}M ({m["direct_pct"]:.1f}%)'],
        ['Indirect Spend', f'EUR {m["indirect_spend"]/1e6:.1f}M ({m["indirect_pct"]:.1f}%)'],
        ['Debit Transactions', f'{m["debit_count"]:,}'],
        ['Credit Transactions', f'{m["credit_count"]:,}'],
        ['Unique Vendors', f'{m.get("total_vendors", 0):,}'],
        ['GL Account Coverage', f'{m.get("gl_pct", 0):.1f}%'],
        ['Profit Center Coverage', f'{m.get("prctr_pct", 0):.1f}%'],
        ['Cost Center Coverage', f'{m.get("kostl_pct", 0):.1f}%'],
    ]
    add_table(doc, ['Metric', 'Value'], rows)
    doc.add_page_break()

def section_sap_tables(doc):
    doc.add_heading('SAP Raw Tables Used', level=1)
    rows = [[k, v['desc'], v['role']] for k, v in SAP_TABLES.items()]
    add_table(doc, ['Table', 'Description', 'Role'], rows)
    doc.add_page_break()

def section_join_logic(doc, mode='rseg'):
    doc.add_heading('Join Logic & Sequence', level=1)
    if mode in ('rseg', 'full'):
        doc.add_paragraph('All joins use LEFT JOIN from RSEG as base table.')
        code = '\n'.join([f'# {j["step"]}. {j["tables"]}  on {j["keys"]} ({j["type"]})' for j in RSEG_JOIN_SEQUENCE])
        add_code_block(doc, code)
    else:
        doc.add_paragraph('FI document analysis uses BKPF headers joined to BSEG line items.')
        add_code_block(doc, '1. BKPF: Filter BUKRS + BLART -> get BELNR/GJAHR\n2. BSEG: Match on BELNR + GJAHR + BUKRS\n3. LFA1: Match on BSEG.LIFNR for vendor name/country')
    doc.add_page_break()

def section_classification(doc, m):
    doc.add_heading('Direct vs Indirect Classification Rules', level=1)
    add_code_block(doc, SPEND_CLASSIFICATION['flowchart'])
    add_table(doc, ['Spend Type', 'Lines', 'Gross EUR', '% of Total'], [
        ['DIRECT', f'{m["direct_lines"]:,}', f'EUR {m["direct_spend"]/1e6:.1f}M', f'{m["direct_pct"]:.1f}%'],
        ['INDIRECT', f'{m["indirect_lines"]:,}', f'EUR {m["indirect_spend"]/1e6:.1f}M', f'{m["indirect_pct"]:.1f}%'],
        ['TOTAL', f'{m["total_records"]:,}', f'EUR {m["total_gross"]/1e6:.1f}M', '100%'],
    ])
    doc.add_page_break()

def section_debit_credit(doc, m):
    doc.add_heading('Debit/Credit (SHKZG) Rules & Analysis', level=1)
    add_table(doc, ['SHKZG', 'Meaning', 'Amount Treatment'], [
        ['S (Soll)', 'Debit / Expense', 'Adds to Gross Spend'],
        ['H (Haben)', 'Credit / Reversal', 'Reduces Net Spend'],
    ])
    add_table(doc, ['Metric', 'Value'], [
        ['Debit Transactions', f'{m["debit_count"]:,}'],
        ['Credit Transactions', f'{m["credit_count"]:,}'],
        ['Gross Spend (Debits)', f'EUR {m["debit_spend"]/1e6:.1f}M'],
        ['Credits/Reversals', f'EUR {m["total_credit"]/1e6:.1f}M'],
        ['Net Spend', f'EUR {m["total_net"]/1e6:.1f}M'],
    ])
    doc.add_page_break()

def section_country_breakdown(doc, by_country, total_gross):
    doc.add_heading('Spend by Company / Country', level=1)
    rows = []
    for _, r in by_country.iterrows():
        pct = r['Gross_EUR'] / total_gross * 100 if total_gross > 0 else 0
        rows.append([str(r['Country']), f'EUR {r["Gross_EUR"]/1e6:.1f}M', f'EUR {r["Net_EUR"]/1e6:.1f}M', f'{int(r["Lines"]):,}', f'{pct:.1f}%'])
    add_table(doc, ['Country', 'Gross EUR', 'Net EUR', 'Lines', '% Total'], rows)
    doc.add_page_break()

def section_gl_analysis(doc, m):
    doc.add_heading('GL Account Analysis', level=1)
    gl = m.get('gl_analysis', pd.DataFrame())
    if len(gl) > 0:
        rows = []
        for _, row in gl.iterrows():
            gs = str(int(row['GL_Account'])) if pd.notna(row['GL_Account']) else 'N/A'
            desc = GL_ACCOUNT_DESCRIPTIONS.get(int(float(row['GL_Account'])), '') if pd.notna(row['GL_Account']) else ''
            rows.append([gs, desc, f'EUR {row["Total_EUR"]/1e6:.1f}M', f'{int(row["Lines"]):,}'])
        add_table(doc, ['GL Account', 'Description', 'Total EUR', 'Lines'], rows)
    doc.add_page_break()

def section_vendor_analysis(doc, m):
    doc.add_heading('Vendor / Supplier Analysis', level=1)
    add_table(doc, ['Metric', 'Value'], [
        ['Total Unique Vendors', f'{m.get("total_vendors", 0):,}'],
        ['Vendor Coverage', f'{m.get("vendor_coverage", 0):.1f}%'],
        ['Top 10 Vendors Share', f'{m.get("top10_pct", 0):.1f}%'],
        ['Top 20 Vendors Share', f'{m.get("top20_pct", 0):.1f}%'],
        ['Pareto 80%', f'{m.get("pareto_80", 0):,} vendors'],
    ])
    vs = m.get('vendor_spend', pd.DataFrame())
    if len(vs) > 0:
        doc.add_heading('Top 20 Vendors', level=2)
        rows = []
        total = m.get('total_dmbtr', m.get('total_gross', 1))
        for i, (_, row) in enumerate(vs.head(20).iterrows(), 1):
            amt = row.get('Total_EUR', row.get('Gross_EUR', 0))
            pct = amt / total * 100 if total > 0 else 0
            rows.append([str(i), safe_vendor_name(row['Vendor'], 40), f'EUR {amt/1e6:.1f}M', f'{int(row["Lines"]):,}', f'{pct:.1f}%'])
        add_table(doc, ['Rank', 'Vendor', 'Total EUR', 'Lines', '% Total'], rows)
    doc.add_page_break()

def section_maverick(doc, m, mav, mav_vendor_spend, high_credit, mode='full'):
    doc.add_heading('Maverick Spend Analysis', level=1)
    tg = mav['total_gross']
    prim_pct = mav['primary']['spend'] / tg * 100 if tg > 0 else 0
    crit_pct = mav['critical']['spend'] / tg * 100 if tg > 0 else 0

    doc.add_heading('Maverick Definition', level=2)
    for tier, info in MAVERICK_TIERS.items():
        doc.add_paragraph(f'{tier}: {info["condition"]} - {info["description"]}')

    doc.add_heading('Maverick Scorecard', level=2)
    rows = [
        ['PRIMARY (No PR+GR)', f'{mav["primary"]["lines"]:,}', f'EUR {mav["primary"]["spend"]/1e6:.1f}M', f'{prim_pct:.1f}%'],
        ['CRITICAL (No PR+GR+Mat)', f'{mav["critical"]["lines"]:,}', f'EUR {mav["critical"]["spend"]/1e6:.1f}M', f'{crit_pct:.1f}%'],
    ]
    if 'three_way_match' in mav:
        tw = mav['three_way_match']
        tw_pct = tw['spend'] / tg * 100 if tg > 0 else 0
        rows.append(['3-Way Match', f'{tw["lines"]:,}', f'EUR {tw["spend"]/1e6:.1f}M', f'{tw_pct:.1f}%'])
    add_table(doc, ['Category', 'Lines', 'Gross EUR', '% Total'], rows)

    if len(mav_vendor_spend) > 0:
        doc.add_heading('Top 20 Maverick Vendors', level=2)
        mv_rows = []
        for i, (_, row) in enumerate(mav_vendor_spend.head(20).iterrows(), 1):
            mv_rows.append([str(i), safe_vendor_name(row['Vendor'], 40),
                            f'EUR {row["Gross_EUR"]/1e6:.2f}M', f'{int(row["Lines"]):,}'])
        add_table(doc, ['Rank', 'Vendor', 'Maverick EUR', 'Lines'], mv_rows)

    if len(high_credit) > 0:
        doc.add_heading('Credit-Heavy Vendors (>20% Ratio)', level=2)
        debit_col = 'Debit_EUR' if 'Debit_EUR' in high_credit.columns else 'Debit_Total'
        credit_col = 'Credit_EUR' if 'Credit_EUR' in high_credit.columns else 'Credit_Total'
        hc_rows = []
        for _, row in high_credit.head(15).iterrows():
            hc_rows.append([safe_vendor_name(row['Vendor'], 40),
                            f'EUR {row[debit_col]/1e6:.2f}M', f'EUR {row[credit_col]/1e6:.2f}M',
                            f'{row["Credit_Ratio"]:.0f}%'])
        add_table(doc, ['Vendor', 'Debit EUR', 'Credit EUR', 'Credit Ratio'], hc_rows)

    doc.add_page_break()

def section_savings_opportunities(doc, m, mav, high_credit, ss_info):
    doc.add_heading('Savings Opportunity Assessment', level=1)
    mav_spend = mav['primary']['spend']
    tg = m['total_gross']; ins = m['indirect_spend']
    hc_credits = high_credit['Credit_EUR'].sum() if 'Credit_EUR' in high_credit.columns else high_credit['Credit_Total'].sum() if len(high_credit)>0 else 0
    ss_spend = ss_info.get('spend', 0)

    rows = [
        ['Maverick Compliance', f'EUR {mav_spend/1e6:.0f}M', '5% of maverick', f'EUR {mav_spend*0.05/1e6:.0f}M'],
        ['Vendor Consolidation', f'EUR {tg/1e6:.0f}M', '3% of total', f'EUR {tg*0.03/1e6:.0f}M'],
        ['Credit Recovery', f'EUR {hc_credits/1e6:.0f}M', 'Investigation needed', 'TBD'],
        ['Dual-Source Leverage', f'EUR {ss_spend/1e6:.0f}M', '2% via competition', f'EUR {ss_spend*0.02/1e6:.0f}M'],
        ['Indirect Optimization', f'EUR {ins/1e6:.0f}M', '7% category mgmt', f'EUR {ins*0.07/1e6:.0f}M'],
    ]
    add_table(doc, ['Opportunity', 'Addressable', 'Estimated Rate', 'Est. Savings'], rows)
    doc.add_page_break()

def section_appendix(doc, m, mav, title="Full Dataset", code="All"):
    doc.add_heading('Appendix: Complete Metrics Summary', level=1)
    prim_pct = mav['primary']['spend'] / mav['total_gross'] * 100 if mav['total_gross'] > 0 else 0
    add_table(doc, ['Metric', 'Value'], [
        ['Analysis Scope', title],
        ['Company Code (BUKRS)', code],
        ['Total Line Items', f'{m["total_records"]:,}'],
        ['Total Gross Spend', f'EUR {m["total_gross"]/1e6:.1f}M'],
        ['Total Net Spend', f'EUR {m["total_net"]/1e6:.1f}M'],
        ['Total Credits', f'EUR {m["total_credit"]/1e6:.1f}M'],
        ['Direct Spend', f'EUR {m["direct_spend"]/1e6:.1f}M ({m["direct_pct"]:.1f}%)'],
        ['Indirect Spend', f'EUR {m["indirect_spend"]/1e6:.1f}M ({m["indirect_pct"]:.1f}%)'],
        ['Unique Vendors', f'{m.get("total_vendors", 0):,}'],
        ['GL Account Coverage', f'{m.get("gl_pct", 0):.1f}%'],
        ['Maverick (Primary)', f'EUR {mav["primary"]["spend"]/1e6:.1f}M ({prim_pct:.1f}%)'],
        ['Analysis Date', DATE],
    ])
